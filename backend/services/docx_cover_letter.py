"""
DOCX Cover Letter Generation Service

Generates professionally formatted DOCX cover letters matching Benjamin Black's
template formatting from the cover_letter_style_guide.md specifications.
"""

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schemas.cover_letter import GeneratedCoverLetter


logger = logging.getLogger(__name__)

# Font constants (matching the Benjamin Black cover letter template)
FONT_NAME = "Georgia"
FONT_NAME_BULLET_CHAR = "Gungsuh"  # Korean font for bullet character (matches resume)
FONT_SIZE_NAME = Pt(16)
FONT_SIZE_CONTACT = Pt(10.5)
FONT_SIZE_BULLET_CHAR = Pt(5.5)  # Smaller bullet character (matches resume)
FONT_SIZE_DATE = Pt(11)
FONT_SIZE_BODY = Pt(11)  # Body text at 11pt for concise appearance

# Line spacing in EMUs (English Metric Units)
# 177800 EMUs = ~12pt single spacing
LINE_SPACING_BODY = Twips(240)  # 12pt
LINE_SPACING_HEADER = Twips(254)  # 12.7pt (from template 203200 EMUs)

# Spacing constants
SPACE_AFTER_NAME = Pt(2)
SPACE_AFTER_CONTACT = Pt(0)
SPACE_BEFORE_DATE = Pt(12)  # Space after horizontal line
SPACE_AFTER_DATE = Pt(12)   # Space before recipient
SPACE_AFTER_RECIPIENT = Pt(0)  # No space between recipient lines
SPACE_AFTER_GREETING = Pt(12)  # Space after "Dear X," before body
SPACE_AFTER_BODY = Pt(12)  # Space between body paragraphs
SPACE_BEFORE_CLOSING = Pt(12)  # Space before "Sincerely,"

# Bullet formatting
BULLET_INDENT = Inches(0.2)  # Hanging indent for bullet text (matches bullet + space width)


def _set_run_font(run, font_name: str, font_size, bold: bool = False, italic: bool = False):
    """Apply font formatting to a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic


def _is_bullet_line(text: str) -> bool:
    """Check if a line is a bullet point."""
    stripped = text.strip()
    return stripped.startswith('- ') or stripped.startswith('• ') or stripped.startswith('* ')


def _ensure_period(text: str) -> str:
    """Ensure text ends with a period (for bullets)."""
    stripped = text.rstrip()
    if stripped and stripped[-1] not in '.!?':
        return stripped + '.'
    return stripped


def _format_bullet_paragraph(para, text: str, font_name: str, font_size):
    """
    Format a bullet point with proper hanging indent.

    The bullet character stays at the left margin, and subsequent lines
    indent to align with the text after the bullet.
    """
    # Strip existing bullet prefix and whitespace
    stripped = text.strip()
    if stripped.startswith('- '):
        stripped = stripped[2:]
    elif stripped.startswith('• '):
        stripped = stripped[2:]
    elif stripped.startswith('* '):
        stripped = stripped[2:]

    # Ensure bullet ends with period
    stripped = _ensure_period(stripped)

    # Set hanging indent: left_indent moves everything right,
    # first_line_indent (negative) pulls first line back to margin
    para.paragraph_format.left_indent = BULLET_INDENT
    para.paragraph_format.first_line_indent = -BULLET_INDENT
    para.paragraph_format.line_spacing = 1.0  # Single line spacing within bullet

    # Add bullet character + tab, then the text
    run = para.add_run("•\t" + stripped)
    _set_run_font(run, font_name, font_size)


def _add_horizontal_line(doc: Document, thickness: float = 0.5, space_before: int = 4, space_after: int = 4):
    """
    Add a thin horizontal line to the document.

    Args:
        doc: The document to add the line to
        thickness: Line thickness in points (default 0.5pt = thin line)
        space_before: Space before line in points (default 4)
        space_after: Space after line in points (default 4)
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)

    # Create a bottom border on the paragraph
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(thickness * 8)))  # Size in 8ths of a point
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    return para


def _format_date(date_str: Optional[str] = None) -> str:
    """Format date for cover letter (e.g., 'December 3, 2025').

    Uses platform-agnostic formatting (works on Windows, Linux, macOS).
    """
    if date_str:
        try:
            dt = datetime.fromisoformat(date_str)
        except ValueError:
            dt = datetime.now()
    else:
        dt = datetime.now()
    # Platform-agnostic: strftime %d is zero-padded, so use day directly
    day = dt.day
    return f"{dt.strftime('%B')} {day}, {dt.year}"


def create_cover_letter_docx(
    cover_letter: GeneratedCoverLetter,
    user_name: str,
    user_email: str,
    user_phone: Optional[str] = None,
    user_linkedin: Optional[str] = None,
    recipient_name: Optional[str] = None,
    date_str: Optional[str] = None
) -> bytes:
    """
    Generate a DOCX cover letter from a GeneratedCoverLetter object.

    Creates a new document from scratch with formatting matching Benjamin Black's
    cover letter template:
    - Georgia font throughout
    - Centered header with name (16pt bold) and contact (10.5pt)
    - Right-aligned date (11pt)
    - Left-aligned recipient block
    - Body paragraphs with proper spacing
    - Closing with signature

    Args:
        cover_letter: The GeneratedCoverLetter object with draft_cover_letter text
        user_name: Full name for the header
        user_email: Email address for the header
        user_phone: Phone number for the header (optional)
        user_linkedin: LinkedIn URL or handle for the header (optional)
        recipient_name: Name of recipient (default: "Hiring Team")
        date_str: ISO date string (default: current date)

    Returns:
        bytes: The generated DOCX file as bytes

    Raises:
        ValueError: If cover letter body text is empty

    Example:
        >>> cl = GeneratedCoverLetter(...)
        >>> docx_bytes = create_cover_letter_docx(
        ...     cover_letter=cl,
        ...     user_name="Benjamin Black",
        ...     user_email="benjamin.black@sloan.mit.edu",
        ...     user_phone="617-504-5529",
        ...     user_linkedin="linkedin.com/in/benjaminblack"
        ... )
    """
    doc = Document()

    # Set up page margins to match template
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.81)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.56)

    # === HEADER SECTION (matching resume header format) ===
    # Name (centered, bold, 16pt)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(4)  # Add space under name (matches resume)
    name_run = name_para.add_run(user_name.upper())
    _set_run_font(name_run, FONT_NAME, FONT_SIZE_NAME, bold=True)

    # Contact line (centered, 10.5pt) - matching resume format with Gungsuh bullets
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = SPACE_AFTER_CONTACT

    # Build contact parts - bullets between items (matching resume)
    contact_parts = [user_email]
    if user_phone:
        contact_parts.append(user_phone)
    if user_linkedin:
        contact_parts.append(user_linkedin)

    # Build contact line with mixed fonts: text in Georgia, bullets in Gungsuh (matches resume)
    for i, part in enumerate(contact_parts):
        if i > 0:
            # Add bullet separator with Gungsuh font (extra spacing to spread content)
            bullet_sep = contact_para.add_run("       ●       ")
            _set_run_font(bullet_sep, FONT_NAME_BULLET_CHAR, FONT_SIZE_BULLET_CHAR)
        # Add contact text in Georgia
        text_run = contact_para.add_run(part)
        _set_run_font(text_run, FONT_NAME, FONT_SIZE_CONTACT)

    # Projects URL line (centered, 10.5pt) with bottom border - no bullet before it (matches resume)
    projects_para = doc.add_paragraph()
    projects_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    projects_para.paragraph_format.space_after = Pt(4)  # Small space after line before date
    # Just spaces, no bullet before portfolio (matches resume)
    space_run = projects_para.add_run("    ")
    _set_run_font(space_run, FONT_NAME, FONT_SIZE_CONTACT)
    projects_run = projects_para.add_run("benjaminblack.consulting/projects")
    _set_run_font(projects_run, FONT_NAME, FONT_SIZE_CONTACT)

    # Add bottom border directly to projects paragraph (like resume header)
    pPr = projects_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # 0.5pt line
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # === DATE SECTION ===
    # Date (right-aligned, 11pt)
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.paragraph_format.space_before = SPACE_BEFORE_DATE
    date_para.paragraph_format.space_after = SPACE_AFTER_DATE

    formatted_date = _format_date(date_str)
    date_run = date_para.add_run(formatted_date)
    _set_run_font(date_run, FONT_NAME, FONT_SIZE_DATE)

    # === RECIPIENT SECTION ===
    # Recipient name (left-aligned, default body size)
    recipient = recipient_name or "Hiring Team"
    recipient_para = doc.add_paragraph()
    recipient_para.paragraph_format.space_after = SPACE_AFTER_RECIPIENT
    recipient_run = recipient_para.add_run(recipient)
    _set_run_font(recipient_run, FONT_NAME, FONT_SIZE_BODY)

    # Company name (if available)
    company_name = cover_letter.company_name
    if company_name:
        company_para = doc.add_paragraph()
        company_para.paragraph_format.space_after = SPACE_AFTER_RECIPIENT
        company_run = company_para.add_run(company_name)
        _set_run_font(company_run, FONT_NAME, FONT_SIZE_BODY)

    # === GREETING ===
    greeting_para = doc.add_paragraph()
    greeting_para.paragraph_format.space_before = SPACE_AFTER_DATE  # Space before greeting
    greeting_para.paragraph_format.space_after = SPACE_AFTER_GREETING  # Space after greeting

    greeting_text = f"Dear {recipient},"
    greeting_run = greeting_para.add_run(greeting_text)
    _set_run_font(greeting_run, FONT_NAME, FONT_SIZE_BODY)

    # === BODY PARAGRAPHS ===
    # Split cover letter text into paragraphs
    body_text = cover_letter.draft_cover_letter

    # Validate non-empty body
    if not body_text or not body_text.strip():
        raise ValueError("Cover letter body text cannot be empty")

    body_paragraphs = [p.strip() for p in body_text.split('\n\n') if p.strip()]

    # If no double-newlines, try single newlines
    if len(body_paragraphs) <= 1:
        body_paragraphs = [p.strip() for p in body_text.split('\n') if p.strip()]

    for i, para_text in enumerate(body_paragraphs):
        # Check if this paragraph contains bullets (lines starting with - or •)
        lines = para_text.split('\n')
        has_bullets = any(_is_bullet_line(line) for line in lines)

        if has_bullets:
            # Process paragraph with mixed content (intro text + bullets)
            for j, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue

                body_para = doc.add_paragraph()

                if _is_bullet_line(line):
                    # Format as bullet with hanging indent
                    body_para.paragraph_format.space_after = Pt(0)  # Minimal spacing between bullets
                    body_para.paragraph_format.space_before = Pt(0)
                    _format_bullet_paragraph(body_para, line, FONT_NAME, FONT_SIZE_BODY)
                else:
                    # Regular text (intro to bullet section)
                    body_para.paragraph_format.space_after = Pt(6)  # Small space before bullets
                    body_run = body_para.add_run(line)
                    _set_run_font(body_run, FONT_NAME, FONT_SIZE_BODY)

            # Add spacing after the bullet section (before next paragraph)
            if i < len(body_paragraphs) - 1:
                # Set larger space after last bullet
                body_para.paragraph_format.space_after = SPACE_AFTER_BODY
        else:
            # Regular paragraph (no bullets)
            body_para = doc.add_paragraph()
            body_para.paragraph_format.space_after = SPACE_AFTER_BODY

            body_run = body_para.add_run(para_text)
            _set_run_font(body_run, FONT_NAME, FONT_SIZE_BODY)

    # === CLOSING ===
    closing_para = doc.add_paragraph()
    closing_para.paragraph_format.space_before = Pt(24)  # Extra line space before closing
    closing_para.paragraph_format.space_after = Pt(0)  # No space after "Sincerely,"
    closing_run = closing_para.add_run("Sincerely,")
    _set_run_font(closing_run, FONT_NAME, FONT_SIZE_BODY)

    # Signature (user name)
    sig_para = doc.add_paragraph()
    sig_para.paragraph_format.space_after = Pt(0)
    sig_run = sig_para.add_run(user_name)
    _set_run_font(sig_run, FONT_NAME, FONT_SIZE_BODY)

    # Save to bytes
    buffer = io.BytesIO()
    try:
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()
    except Exception as e:
        logger.error(f"Failed to serialize DOCX document: {str(e)}", exc_info=True)
        raise ValueError(f"Failed to generate DOCX document: {str(e)}")

    logger.info(f"Generated DOCX cover letter for {cover_letter.company_name or 'unknown company'}")
    return docx_bytes


def create_cover_letter_docx_from_dict(
    cover_letter_json: dict,
    user_name: str,
    user_email: str,
    user_phone: Optional[str] = None,
    user_linkedin: Optional[str] = None,
    recipient_name: Optional[str] = None,
    date_str: Optional[str] = None
) -> bytes:
    """
    Generate a DOCX cover letter from a dictionary (JSON) representation.

    Convenience wrapper for create_cover_letter_docx that accepts a dict
    instead of a GeneratedCoverLetter object.

    Args:
        cover_letter_json: Dictionary with at minimum 'draft_cover_letter' key
        user_name: Full name for the header
        user_email: Email address for the header
        user_phone: Phone number for the header (optional)
        user_linkedin: LinkedIn URL or handle for the header (optional)
        recipient_name: Name of recipient (default: "Hiring Team")
        date_str: ISO date string (default: current date)

    Returns:
        bytes: The generated DOCX file as bytes
    """
    # Create a minimal GeneratedCoverLetter-like object
    class CoverLetterData:
        def __init__(self, data: dict):
            self.draft_cover_letter = data.get("draft_cover_letter", "")
            self.company_name = data.get("company_name")
            self.job_title = data.get("job_title")

    cl_data = CoverLetterData(cover_letter_json)

    return create_cover_letter_docx(
        cover_letter=cl_data,
        user_name=user_name,
        user_email=user_email,
        user_phone=user_phone,
        user_linkedin=user_linkedin,
        recipient_name=recipient_name,
        date_str=date_str
    )
