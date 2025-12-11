"""
DOCX Resume Generation Service

Generates professionally formatted DOCX resumes from TailoredResume JSON.
Matches the formatting of the Benjamin Black resume template.

v1.3.0: Added engagement-aware rendering for consulting roles
"""

import io
import logging
from typing import Optional, List

from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schemas.resume_tailor import (
    TailoredResume, SelectedRole, SelectedSkill, SelectedBullet, SelectedEngagement
)

# Import constants from config
try:
    from config.constants import BBC_CLIENTS
except ImportError:
    # Fallback for backward compatibility
    BBC_CLIENTS = [
        {"name": "Edward Jones", "start": "2/2023", "end": "11/2023"},
        {"name": "Darling Consulting Group", "start": "9/2022", "end": "1/2023"},
        {"name": "Squark (Machine Learning / AI Startup)", "start": "5/2018", "end": "7/2020"},
        {"name": "Vestmark, Inc.", "start": "7/2017", "end": "3/2018"},
        {"name": "John Hancock Investments (through Olmstead Associates)", "start": "11/2016", "end": "1/2017"},
        {"name": "Olmstead Associates", "start": "1/2017", "end": "5/2017"},
        {"name": "Fidelity Investments", "start": "8/2016", "end": "11/2016"},
    ]

# Approximate lines per page for multi-page detection (conservative estimate)
LINES_PER_PAGE = 45
LINES_PER_BULLET = 2  # Average lines per bullet including spacing


logger = logging.getLogger(__name__)

# Font constants (matching the original resume)
FONT_NAME = "Georgia"
FONT_NAME_BULLET_CHAR = "Gungsuh"  # Korean font for bullet character
FONT_SIZE_NAME = Pt(16)
FONT_SIZE_CONTACT = Pt(10.5)
FONT_SIZE_SECTION_HEADER = Pt(11)
FONT_SIZE_COMPANY = Pt(11.5)
FONT_SIZE_COMPANY_DETAIL = Pt(10)
FONT_SIZE_JOB_TITLE = Pt(10)
FONT_SIZE_BULLET = Pt(10.5)
FONT_SIZE_BULLET_CHAR = Pt(5.5)  # Smaller bullet character
FONT_SIZE_SKILL_CATEGORY = Pt(10.5)

# Indentation constants
INDENT_SUMMARY = Inches(0.125)
INDENT_SECTION = Inches(0.13)
# Bullet spacing constant - distance from bullet to text start
BULLET_TEXT_GAP = Inches(0.18)  # Space between bullet and text

# Role bullet indentation: bullet in LEFT MARGIN (negative), text at 0" (aligned with company)
# first_line_indent pulls the bullet into the margin, left_indent keeps wrapped text at 0"
INDENT_BULLET_LEFT = Inches(0.0)  # Text wraps at 0" (aligned with company name)
INDENT_BULLET_HANGING = -BULLET_TEXT_GAP  # Bullet extends into margin

# Engagement bullet indentation: bullet at 0" (company level), text at 0.18" from bullet
INDENT_ENGAGEMENT_BULLET_LEFT = BULLET_TEXT_GAP  # Text wraps at 0.18"
INDENT_ENGAGEMENT_BULLET_HANGING = -BULLET_TEXT_GAP  # First line starts at 0"

# Tab stop for right-aligned dates (at 6" from left margin = 7.5" - 0.5" right margin - 1" buffer)
DATE_TAB_POSITION = Inches(7.0)  # 6" from 0.5" left margin = 6.5" total, but page is 7.5" wide

# Spacing constants
SPACE_AFTER_SUMMARY = Pt(0)  # Reduced - horizontal line provides spacing
SPACE_BEFORE_SECTION = Pt(10)
SPACE_AFTER_SECTION = Pt(6)
SPACE_BEFORE_COMPANY = Pt(7)
SPACE_AFTER_COMPANY = Pt(2)
SPACE_AFTER_TITLE = Pt(2)
SPACE_AFTER_BULLET = Pt(2)


def _add_horizontal_line(doc: Document, width: float = 7.5, thickness: float = 0.5, space_after: float = 2):
    """
    Add a thin horizontal line to the document.

    Args:
        doc: The document to add the line to
        width: Width of the line in inches (default 7.5 = full width with margins)
        thickness: Line thickness in points (default 0.5pt = thin line)
        space_after: Space after the line in points (default 2pt)
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(space_after)

    # Add a space character to establish the font (empty string doesn't work)
    # At 0.5pt the space is virtually invisible but establishes Georgia font
    tiny_run = para.add_run(" ")
    _set_run_font(tiny_run, FONT_NAME, Pt(0.5))

    # Set line spacing to Exactly 8pt to minimize vertical space
    para.paragraph_format.line_spacing = Pt(8)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    # Create a bottom border on the paragraph
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(thickness * 8)))  # Size in 8ths of a point
    bottom.set(qn('w:space'), '0')  # No space between text and line
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    return para


def _add_section_header(doc: Document, title: str):
    """
    Add a section header with horizontal line above and underline below.

    Creates properly formatted section headers (Professional Experience, Technical Skills, Education)
    with small caps styling and underline border.

    Args:
        doc: The document to add to
        title: Section title text (e.g., "Professional Experience")

    Returns:
        The header paragraph
    """
    # Add line above section
    _add_horizontal_line(doc, space_after=0)

    # Section header paragraph
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(6)

    # Add the title text with small caps
    header_run = header_para.add_run(title)
    _set_run_font(header_run, FONT_NAME, FONT_SIZE_SECTION_HEADER, bold=True, small_caps=True)

    # Add bottom border (underline effect)
    pPr = header_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    return header_para


def _add_header_line(header):
    """
    Add a thin horizontal line to the document header, tight to contact info with 12pt space below.

    Args:
        header: The header section to add the line to
    """
    para = header.add_paragraph()
    para.paragraph_format.space_before = Pt(0)  # No space above - tight to contact info
    para.paragraph_format.space_after = Pt(12)  # 12pt space under the line

    # Create a bottom border on the paragraph
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # 0.5pt (4/8)
    bottom.set(qn('w:space'), '0')  # No space between text and line
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    return para


def _add_right_tab_stop(paragraph, position=DATE_TAB_POSITION):
    """
    Add a right-aligned tab stop to the paragraph for date alignment.

    Args:
        paragraph: The paragraph to add the tab stop to
        position: Position of the tab stop from left margin (default: DATE_TAB_POSITION)
    """
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(position, WD_TAB_ALIGNMENT.RIGHT)


def _set_keep_together(paragraph):
    """Set the Keep Together property on a paragraph to prevent page breaks within."""
    pPr = paragraph._p.get_or_add_pPr()
    keepLines = OxmlElement('w:keepLines')
    keepLines.set(qn('w:val'), 'true')
    pPr.append(keepLines)


def _set_keep_with_next(paragraph):
    """Set the Keep With Next property on a paragraph to prevent page breaks after."""
    pPr = paragraph._p.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    keepNext.set(qn('w:val'), 'true')
    pPr.append(keepNext)


def _add_page_break(doc: Document):
    """Add a page break to the document."""
    from docx.enum.text import WD_BREAK
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _set_run_font(run, font_name: str, font_size, bold: bool = False,
                  italic: bool = False, underline: bool = False, small_caps: bool = False):
    """Apply font formatting to a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.small_caps = small_caps


def _add_bullet_point(doc: Document, text: str, indent_left: float = INDENT_BULLET_LEFT,
                      indent_hanging: float = INDENT_BULLET_HANGING):
    """Add a bullet point paragraph with proper formatting.

    Uses left_indent to set where wrapped text aligns, and first_line_indent (negative)
    to pull the bullet character back to its proper position.
    """
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = indent_left
    para.paragraph_format.first_line_indent = indent_hanging
    para.paragraph_format.space_after = SPACE_AFTER_BULLET

    # Add tab stop for bullet-to-text spacing (at the left_indent position)
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(indent_left, WD_TAB_ALIGNMENT.LEFT)

    # Add bullet character (Gungsuh 5.5pt for proper size)
    bullet_run = para.add_run("●\t")
    _set_run_font(bullet_run, FONT_NAME_BULLET_CHAR, FONT_SIZE_BULLET_CHAR)

    # Add text
    text_run = para.add_run(text)
    _set_run_font(text_run, FONT_NAME, FONT_SIZE_BULLET)

    return para


def _format_date_range(start_date: str, end_date: Optional[str]) -> str:
    """Format date range for display (e.g., '9/2024 – Present')."""
    def format_date(date_str: str) -> str:
        if not date_str:
            return "Present"
        parts = date_str.split("-")
        if len(parts) >= 2:
            month = int(parts[1])
            year = parts[0]
            return f"{month}/{year}"
        return date_str

    start = format_date(start_date)
    end = format_date(end_date) if end_date else "Present"
    return f"{start} – {end}"


def _parse_company_name_with_parenthetical(name: str) -> tuple[str, str | None]:
    """
    Parse company name to extract parenthetical descriptor.

    Examples:
        "KeyLogic Associates (assigned to Kessel Run)" -> ("KeyLogic Associates", "assigned to Kessel Run")
        "Capital Group | Los Angeles, CA" -> ("Capital Group | Los Angeles, CA", None)
        "ACME Corp" -> ("ACME Corp", None)

    Returns:
        Tuple of (base_name, parenthetical_text or None)
    """
    import re
    # Match pattern: "Company Name (parenthetical text)" at end of string
    # But exclude simple location patterns like "(Boston, MA)" which are handled separately
    match = re.match(r'^(.+?)\s+\(([^)]+)\)\s*$', name)
    if match:
        base_name = match.group(1)
        paren_text = match.group(2)
        # Don't split if it looks like a location (City, ST format)
        if re.match(r'^[A-Z][a-z]+,\s*[A-Z]{2}$', paren_text):
            return (name, None)
        return (base_name, paren_text)
    return (name, None)


def _add_company_name_with_parenthetical(para, company_name: str, font_size=FONT_SIZE_COMPANY):
    """
    Add company name to paragraph, with parenthetical text in regular (not bold) italics.

    The base company name is bold+underlined, parenthetical descriptors are italic only.
    Example: "KeyLogic Associates (assigned to Kessel Run)"
             -> KeyLogic Associates [bold, underlined] (assigned to Kessel Run) [italic, not bold]

    Args:
        para: The paragraph to add runs to
        company_name: Full company name possibly with parenthetical
        font_size: Font size to use
    """
    base_name, paren_text = _parse_company_name_with_parenthetical(company_name)

    # Add base company name (bold, underlined)
    company_run = para.add_run(base_name)
    _set_run_font(company_run, FONT_NAME, font_size, bold=True, underline=True)

    # Add parenthetical text if present (Georgia 10pt italic, explicitly NOT bold)
    if paren_text:
        paren_run = para.add_run(f" ({paren_text})")
        _set_run_font(paren_run, FONT_NAME, FONT_SIZE_JOB_TITLE, bold=False, italic=True)


def _is_consulting_role(role: SelectedRole) -> bool:
    """
    Check if this role is a consulting role that should show engagements.

    v1.3.0: Uses employer_type field if available, otherwise falls back to
    name-based detection for backward compatibility.
    """
    # Check employer_type first (v1.3.0)
    if role.employer_type == "independent_consulting":
        return True

    # Check if role has engagements (v1.3.0)
    if role.selected_engagements:
        return True

    # Fallback: name-based detection for legacy data
    employer_lower = role.employer_name.lower() if role.employer_name else ""
    return "benjamin black consulting" in employer_lower


def _is_bbc_role(role: SelectedRole) -> bool:
    """
    Check if this role is the Benjamin Black Consulting role.

    NOTE: Deprecated - use _is_consulting_role instead.
    """
    return _is_consulting_role(role)


def _get_client_bullets(bullets: list[SelectedBullet], client_name: str) -> list[SelectedBullet]:
    """
    Get bullets that are associated with a specific BBC client.

    Matches based on client company name appearing in bullet text or tags.
    """
    matched = []
    client_lower = client_name.lower()

    # Extract primary company name (before parentheses) and create search variants
    primary_name = client_name.split("(")[0].strip().lower()

    # Client-specific search keywords to improve matching
    CLIENT_KEYWORDS = {
        "edward jones": ["edward jones"],
        "darling consulting group": ["darling consulting", "dcg"],
        "squark": ["squark", "ml/ai startup"],
        "vestmark": ["vestmark", "broker dealers", "wealth managers"],
        "john hancock": ["john hancock", "sec modernization"],
        "olmstead": ["olmstead"],
        "fidelity": ["fidelity finance", "fidelity"],
    }

    # Get keywords for this client
    keywords = []
    for key, kw_list in CLIENT_KEYWORDS.items():
        if key in primary_name:
            keywords = kw_list
            break
    if not keywords:
        keywords = [primary_name]

    for bullet in bullets:
        bullet_text_lower = bullet.text.lower() if bullet.text else ""
        tags_lower = [t.lower() for t in (bullet.tags or [])]

        # Check if any keyword appears in bullet text or tags
        for keyword in keywords:
            if (keyword in bullet_text_lower or
                any(keyword in tag for tag in tags_lower)):
                matched.append(bullet)
                break  # Found a match, don't add same bullet twice

    return matched


def _add_bbc_client_entry(doc: Document, client_name: str, date_range: str,
                          bullets: list[SelectedBullet], keep_together: bool = True):
    """
    Add a single BBC client entry with bold name, date in parens, and indented bullets.

    Args:
        doc: The document to add to
        client_name: Client company name (e.g., "Edward Jones")
        date_range: Date range string (e.g., "2/2023-11/2023")
        bullets: List of selected bullets for this client
        keep_together: Whether to apply Keep Together formatting
    """
    # Client name line: bold name with date in parentheses - aligned at 0"
    client_para = doc.add_paragraph()
    client_para.paragraph_format.left_indent = Inches(0)  # Align with company name
    client_para.paragraph_format.space_before = Pt(4)
    client_para.paragraph_format.space_after = Pt(2)

    # Bold client name
    name_run = client_para.add_run(client_name)
    _set_run_font(name_run, FONT_NAME, FONT_SIZE_BULLET, bold=True)

    # Date in parentheses (not bold)
    date_run = client_para.add_run(f" ({date_range})")
    _set_run_font(date_run, FONT_NAME, Pt(9.5))

    # Keep client header with first bullet (header should never orphan)
    if bullets:
        _set_keep_with_next(client_para)
    _set_keep_together(client_para)

    # Add bullets for this client (text indented, bullet aligns with company name)
    for i, bullet in enumerate(bullets):
        bullet_para = _add_bullet_point(
            doc, bullet.text,
            indent_left=INDENT_ENGAGEMENT_BULLET_LEFT,  # 0.18" - text indented
            indent_hanging=INDENT_ENGAGEMENT_BULLET_HANGING  # -0.18" - bullet at company name
        )
        # Keep bullets together - each bullet stays with the next one (except last)
        _set_keep_together(bullet_para)
        if i < len(bullets) - 1:
            _set_keep_with_next(bullet_para)


def _add_engagement_entry(doc: Document, engagement: SelectedEngagement, keep_together: bool = True):
    """
    Add a single engagement entry with client name, project, and bullets.

    v1.3.0: New engagement-aware formatting for consulting roles.

    Args:
        doc: The document to add to
        engagement: The SelectedEngagement with client, project, and bullets
        keep_together: Whether to apply Keep Together formatting
    """
    # Engagement header line - aligned at 0" (with company name)
    eng_para = doc.add_paragraph()
    eng_para.paragraph_format.left_indent = Inches(0)  # Align with company name
    eng_para.paragraph_format.space_before = Pt(4)
    eng_para.paragraph_format.space_after = Pt(2)

    # Client name - parse for parenthetical text like "(via Olmstead Associates)"
    client_name = engagement.client or ""
    if client_name:
        base_name, paren_text = _parse_company_name_with_parenthetical(client_name)

        # Base client name in bold
        client_run = eng_para.add_run(base_name)
        _set_run_font(client_run, FONT_NAME, FONT_SIZE_BULLET, bold=True)

        # Parenthetical text in italic (NOT bold)
        if paren_text:
            paren_run = eng_para.add_run(f" ({paren_text})")
            _set_run_font(paren_run, FONT_NAME, FONT_SIZE_BULLET, bold=False, italic=True)

    # Em-dash and project name NOT bold
    if engagement.project_name:
        separator = " — " if client_name else ""
        project_run = eng_para.add_run(f"{separator}{engagement.project_name}")
        _set_run_font(project_run, FONT_NAME, FONT_SIZE_BULLET, bold=False)

    # Note: date_range_label omitted for cleaner engagement display

    # Keep engagement header with first bullet (engagement header should never orphan)
    if engagement.selected_bullets:
        _set_keep_with_next(eng_para)
    _set_keep_together(eng_para)

    # Add bullets for this engagement (text indented, bullet aligns with company name)
    for i, bullet in enumerate(engagement.selected_bullets):
        bullet_para = _add_bullet_point(
            doc, bullet.text,
            indent_left=INDENT_ENGAGEMENT_BULLET_LEFT,  # 0.18" - text indented
            indent_hanging=INDENT_ENGAGEMENT_BULLET_HANGING  # -0.18" - bullet at company name
        )
        # Keep bullets together - each bullet stays with the next one (except last)
        _set_keep_together(bullet_para)
        if i < len(engagement.selected_bullets) - 1:
            _set_keep_with_next(bullet_para)


def _add_continuation_header(doc: Document, role: SelectedRole):
    """
    Add a continuation header for a consulting role on page 2.

    Renders: Company Name (continued) | Location    Date Range
    Without job title (already shown on page 1).

    Args:
        doc: The document to add to
        role: The SelectedRole (consulting type)
    """
    company_para = doc.add_paragraph()
    company_para.paragraph_format.space_before = SPACE_BEFORE_COMPANY
    company_para.paragraph_format.space_after = SPACE_AFTER_COMPANY
    company_para.paragraph_format.line_spacing = 1.17

    # Add right-aligned tab stop for date
    _add_right_tab_stop(company_para)

    # Company name with "(continued)" suffix
    company_text = f"{role.employer_name} (continued)"
    _add_company_name_with_parenthetical(company_para, company_text)

    # Location if present
    if role.location:
        loc_run = company_para.add_run(f" | {role.location}")
        _set_run_font(loc_run, FONT_NAME, FONT_SIZE_COMPANY_DETAIL)

    # Date range
    date_range = _format_date_range(role.start_date, role.end_date)
    company_para.add_run("\t")
    date_run = company_para.add_run(date_range)
    _set_run_font(date_run, FONT_NAME, FONT_SIZE_COMPANY_DETAIL)

    # Keep header with first engagement
    _set_keep_with_next(company_para)
    _set_keep_together(company_para)


def _add_consulting_experience_entry(doc: Document, role: SelectedRole, continued: bool = False):
    """
    Add a consulting role entry with engagements.

    v1.3.0: Uses selected_engagements if available, otherwise falls back to
    legacy BBC_CLIENTS matching for backward compatibility.

    v1.4.0: Supports page_preference on engagements for automatic page breaks
    with continuation headers.

    Args:
        doc: The document to add to
        role: The SelectedRole (consulting type)
        continued: Whether this is a continuation from previous page
    """
    # Company line with date (same as regular entry)
    company_para = doc.add_paragraph()
    company_para.paragraph_format.space_before = SPACE_BEFORE_COMPANY
    company_para.paragraph_format.space_after = SPACE_AFTER_COMPANY
    company_para.paragraph_format.line_spacing = 1.17

    # Add right-aligned tab stop for date
    _add_right_tab_stop(company_para)

    # Company name with parenthetical in italics, with optional "(continued)"
    company_text = role.employer_name
    if continued:
        company_text = f"{company_text} (continued)"
    _add_company_name_with_parenthetical(company_para, company_text)

    # Location if present (separate from parenthetical in company name)
    if role.location:
        loc_run = company_para.add_run(f" | {role.location}")
        _set_run_font(loc_run, FONT_NAME, FONT_SIZE_COMPANY_DETAIL)

    # Date range (using single tab to right-aligned tab stop)
    date_range = _format_date_range(role.start_date, role.end_date)
    company_para.add_run("\t")  # Single tab to right-aligned stop
    date_run = company_para.add_run(date_range)
    _set_run_font(date_run, FONT_NAME, FONT_SIZE_COMPANY_DETAIL)

    # Apply keep together to header
    _set_keep_with_next(company_para)
    _set_keep_together(company_para)

    # Job title (bold, italic) - only on first page (not continued)
    if not continued:
        title_para = doc.add_paragraph()
        title_para.paragraph_format.left_indent = Inches(0)  # Align with company name
        title_para.paragraph_format.space_after = SPACE_AFTER_TITLE

        title_run = title_para.add_run(role.job_title)
        _set_run_font(title_run, FONT_NAME, FONT_SIZE_JOB_TITLE, bold=True, italic=True)
        _set_keep_with_next(title_para)
        _set_keep_together(title_para)

        # Role summary if provided (v1.3.0) - NOT italic
        if role.role_summary:
            summary_para = doc.add_paragraph()
            summary_para.paragraph_format.left_indent = Inches(0)  # Align with company name
            summary_para.paragraph_format.space_after = Pt(4)
            summary_run = summary_para.add_run(role.role_summary)
            _set_run_font(summary_run, FONT_NAME, FONT_SIZE_BULLET, italic=False)  # Not italic
            _set_keep_with_next(summary_para)

    # Check if we have v1.3.0 engagements
    if role.selected_engagements:
        # Track if we've inserted the page break for page 2 engagements
        page_break_inserted = False

        for engagement in role.selected_engagements:
            # Check if this engagement should start on page 2
            # Handle both int and string values (DB may store as either)
            page_pref = engagement.page_preference
            should_break = (page_pref == 2 or page_pref == "2") and not page_break_inserted
            if should_break:
                # Insert page break before first page 2 engagement
                _add_page_break(doc)
                # Add continuation header for page 2
                _add_continuation_header(doc, role)
                page_break_inserted = True

            _add_engagement_entry(doc, engagement, keep_together=True)
    elif role.selected_bullets:
        # No engagements but has direct bullets - render like a normal role
        # This handles BBC periods without client engagements (e.g., current period)
        for bullet in role.selected_bullets:
            _add_bullet_point(doc, bullet.text)


def _add_bbc_experience_entry(doc: Document, role: SelectedRole):
    """
    Add the Benjamin Black Consulting experience entry with all clients listed.

    This special formatting shows all BBC clients in fixed order with their
    associated bullets, using Keep Together formatting to ensure the section
    stays on page 1.

    NOTE: This is the legacy function. For v1.3.0+, use _add_consulting_experience_entry
    which supports the engagement structure.
    """
    # Use the new consulting entry function
    _add_consulting_experience_entry(doc, role, continued=False)


def _group_skills_by_category(skills: list[SelectedSkill]) -> dict[str, list[str]]:
    """
    Group skills by category for display.

    Uses common category prefixes to organize skills logically.
    """
    if not skills:
        return {}

    # Define category mappings based on common skill patterns
    category_patterns = {
        "AI/ML": ["ai", "ml", "machine learning", "deep learning", "nlp", "llm",
                  "rag", "vector", "embedding", "prompt", "neural", "transformer"],
        "Programming": ["python", "sql", "r ", "java", "scala", "spark", "pandas",
                       "numpy", "scikit", "tensorflow", "pytorch", "code", "script"],
        "Data": ["data", "etl", "pipeline", "warehouse", "lake", "hadoop", "kafka",
                "snowflake", "databricks", "dbt", "airflow", "analytics"],
        "Cloud & Infrastructure": ["aws", "azure", "gcp", "cloud", "docker",
                                   "kubernetes", "terraform", "devops", "ci/cd"],
        "Governance & Strategy": ["governance", "strategy", "compliance", "policy",
                                  "framework", "architecture", "leadership", "management"],
        "Visualization & BI": ["tableau", "power bi", "looker", "dashboard",
                              "visualization", "reporting", "qlik"],
    }

    categorized: dict[str, list[str]] = {}
    uncategorized: list[str] = []

    for skill in skills:
        skill_lower = skill.skill.lower()
        matched = False

        for category, patterns in category_patterns.items():
            if any(pattern in skill_lower for pattern in patterns):
                if category not in categorized:
                    categorized[category] = []
                categorized[category].append(skill.skill)
                matched = True
                break

        if not matched:
            uncategorized.append(skill.skill)

    # Add uncategorized skills to "Other" if any
    if uncategorized:
        categorized["Other"] = uncategorized

    return categorized


def create_resume_docx(
    tailored_resume: TailoredResume,
    user_name: Optional[str] = None,
    user_email: Optional[str] = None,
    user_phone: Optional[str] = None,
    user_linkedin: Optional[str] = None,
    user_portfolio: Optional[str] = None,
    education: Optional[list[dict]] = None
) -> bytes:
    """
    Generate a DOCX resume from a TailoredResume object.

    Creates a new document from scratch with formatting matching the
    Benjamin Black resume template.

    Args:
        tailored_resume: The TailoredResume JSON object with selected content
        user_name: Full name for the header
        user_email: Email address for the header
        user_phone: Phone number for the header (optional)
        user_linkedin: LinkedIn URL or handle for the header (optional)
        user_portfolio: Portfolio URL for the header (optional)
        education: List of education entries, each with:
            - institution: str
            - location: str (optional)
            - degree: str
            - details: list[str] (optional bullet points)

    Returns:
        bytes: The generated DOCX file as bytes

    Example:
        >>> resume = TailoredResume(...)
        >>> docx_bytes = create_resume_docx(
        ...     tailored_resume=resume,
        ...     user_name="Benjamin Black",
        ...     user_email="benjamin.black@sloan.mit.edu",
        ...     user_phone="617-504-5529",
        ...     user_linkedin="linkedin.com/in/benjaminblack",
        ...     user_portfolio="benjaminblack.ai"
        ... )
    """
    doc = Document()

    # Set up page margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.56)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # === HEADER SECTION (in document header) ===
    header = section.header

    # Name
    name_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(4)  # Add space under name
    display_name = (user_name or "").upper()
    name_run = name_para.add_run(display_name)
    _set_run_font(name_run, FONT_NAME, FONT_SIZE_NAME, bold=True)

    # Contact line
    contact_para = header.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(12)  # 12pt space after the line

    # Build contact parts - bullets between items, but NO bullet before portfolio (last item)
    contact_parts = []
    if user_email:
        contact_parts.append(user_email)
    if user_phone:
        contact_parts.append(user_phone)
    if user_linkedin:
        contact_parts.append(user_linkedin)

    # Build contact line with mixed fonts: text in Georgia, bullets in Gungsuh
    for i, part in enumerate(contact_parts):
        if i > 0:
            # Add bullet separator with Gungsuh font (extra spacing to spread content)
            bullet_sep = contact_para.add_run("       ●       ")
            _set_run_font(bullet_sep, FONT_NAME_BULLET_CHAR, FONT_SIZE_BULLET_CHAR)
        # Add contact text in Georgia
        text_run = contact_para.add_run(part)
        _set_run_font(text_run, FONT_NAME, FONT_SIZE_CONTACT)

    # Add portfolio at the end WITHOUT a bullet separator (just spaces)
    if user_portfolio:
        space_run = contact_para.add_run("    ")
        _set_run_font(space_run, FONT_NAME, FONT_SIZE_CONTACT)
        portfolio_run = contact_para.add_run(user_portfolio)
        _set_run_font(portfolio_run, FONT_NAME, FONT_SIZE_CONTACT)

    # Add bottom border directly to contact line (underline effect, tight to text)
    pPr = contact_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')  # 0.5pt
    bottom.set(qn('w:space'), '1')  # minimal space between text and line
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # === SUMMARY ===
    if tailored_resume.tailored_summary:
        summary_para = doc.add_paragraph()
        summary_para.paragraph_format.left_indent = INDENT_SUMMARY
        summary_para.paragraph_format.space_after = SPACE_AFTER_SUMMARY
        summary_para.paragraph_format.line_spacing = 1.17  # ~7/6

        # Add summary text
        summary_text = tailored_resume.tailored_summary.rstrip()
        # Ensure it ends with a period before adding MIT Sloan MBA
        if not summary_text.endswith('.'):
            summary_text += '.'
        summary_text += " MIT Sloan MBA."

        summary_run = summary_para.add_run(summary_text)
        _set_run_font(summary_run, FONT_NAME, FONT_SIZE_BULLET)

    # === PROFESSIONAL EXPERIENCE ===
    if tailored_resume.selected_roles:
        _add_section_header(doc, "Professional Experience")

        # Add each role
        for role in tailored_resume.selected_roles:
            if _is_consulting_role(role):
                # Engagement-based formatting for consulting roles
                _add_consulting_experience_entry(doc, role, continued=False)
            else:
                _add_experience_entry(doc, role)

    # === TECHNICAL SKILLS ===
    if tailored_resume.selected_skills:
        _add_section_header(doc, "Technical Skills")

        # Group and add skills
        grouped_skills = _group_skills_by_category(tailored_resume.selected_skills)
        for category, skill_list in grouped_skills.items():
            skills_para = doc.add_paragraph()
            skills_para.paragraph_format.left_indent = Inches(0)  # Align with company names
            skills_para.paragraph_format.space_after = Pt(2)

            # Category name (bold)
            cat_run = skills_para.add_run(f"{category}: ")
            _set_run_font(cat_run, FONT_NAME, FONT_SIZE_SKILL_CATEGORY, bold=True)

            # Skills list
            skills_text = ", ".join(skill_list)
            skills_run = skills_para.add_run(skills_text)
            _set_run_font(skills_run, FONT_NAME, FONT_SIZE_SKILL_CATEGORY)

    # === EDUCATION ===
    if education:
        _add_section_header(doc, "Education")

        for edu_entry in education:
            _add_education_entry(doc, edu_entry)

    # Save to bytes with explicit error handling
    buffer = io.BytesIO()
    try:
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()
    except Exception as e:
        logger.error(f"Failed to serialize DOCX document: {str(e)}", exc_info=True)
        raise ValueError(f"Failed to generate DOCX document: {str(e)}")

    logger.info(f"Generated DOCX resume with {len(tailored_resume.selected_roles)} roles")
    return docx_bytes


def _add_experience_entry(doc: Document, role: SelectedRole):
    """Add a single experience entry with company, title, and bullets."""
    # Company line with date
    company_para = doc.add_paragraph()
    company_para.paragraph_format.space_before = SPACE_BEFORE_COMPANY
    company_para.paragraph_format.space_after = SPACE_AFTER_COMPANY
    company_para.paragraph_format.line_spacing = 1.17

    # Add right-aligned tab stop for date
    _add_right_tab_stop(company_para)

    # Company name with parenthetical in italics (e.g., "KeyLogic (assigned to Kessel Run)")
    _add_company_name_with_parenthetical(company_para, role.employer_name)

    # Location if present (separate from parenthetical in company name)
    if role.location:
        loc_run = company_para.add_run(f" | {role.location}")
        _set_run_font(loc_run, FONT_NAME, FONT_SIZE_COMPANY_DETAIL)

    # Date range (using single tab to right-aligned tab stop)
    date_range = _format_date_range(role.start_date, role.end_date)
    company_para.add_run("\t")  # Single tab to right-aligned stop
    date_run = company_para.add_run(date_range)
    _set_run_font(date_run, FONT_NAME, FONT_SIZE_COMPANY_DETAIL)

    # Job title (bold, italic) - aligned with company name
    title_para = doc.add_paragraph()
    title_para.paragraph_format.left_indent = Inches(0)  # Align with company name
    title_para.paragraph_format.space_after = SPACE_AFTER_TITLE

    title_run = title_para.add_run(role.job_title)
    _set_run_font(title_run, FONT_NAME, FONT_SIZE_JOB_TITLE, bold=True, italic=True)

    # Bullets
    for bullet in role.selected_bullets:
        _add_bullet_point(doc, bullet.text)


def _add_education_entry(doc: Document, edu: dict):
    """Add a single education entry."""
    # Institution line
    inst_para = doc.add_paragraph()
    inst_para.paragraph_format.space_before = SPACE_BEFORE_COMPANY
    inst_para.paragraph_format.space_after = SPACE_AFTER_COMPANY
    inst_para.paragraph_format.line_spacing = 1.17

    # Institution name (bold, underlined)
    inst_run = inst_para.add_run(edu.get("institution", ""))
    _set_run_font(inst_run, FONT_NAME, FONT_SIZE_COMPANY, bold=True, underline=True)

    # Location if present
    location = edu.get("location")
    if location:
        loc_run = inst_para.add_run(f" ({location})")
        _set_run_font(loc_run, FONT_NAME, FONT_SIZE_COMPANY_DETAIL)

    # Degree (bold, italic) - no indentation, aligns with institution name
    degree = edu.get("degree")
    if degree:
        degree_para = doc.add_paragraph()
        degree_para.paragraph_format.left_indent = Inches(0)  # No indentation
        degree_para.paragraph_format.space_after = SPACE_AFTER_TITLE

        degree_run = degree_para.add_run(degree)
        _set_run_font(degree_run, FONT_NAME, FONT_SIZE_JOB_TITLE, bold=True, italic=True)

    # Detail bullets if present
    details = edu.get("details", [])
    for detail in details:
        _add_bullet_point(doc, detail)
